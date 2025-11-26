from typing import List
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import csv
import io
from elasticsearch import Elasticsearch
from app.config import settings

class DocumentExporter:
    """Export search results to various formats"""
    
    def __init__(self):
        self.es = Elasticsearch([f"http://{settings.ELASTICSEARCH_HOST}:{settings.ELASTICSEARCH_PORT}"])
        self.index_name = settings.ELASTICSEARCH_INDEX
    
    def _get_documents(self, doc_ids: List[str]) -> List[dict]:
        """Retrieve documents from Elasticsearch"""
        docs = []
        for doc_id in doc_ids:
            try:
                result = self.es.get(index=self.index_name, id=doc_id)
                docs.append(result['_source'])
            except:
                pass
        return docs
    
    def export_to_pdf(self, doc_ids: List[str], include_summary: bool = False) -> bytes:
        """Export documents to PDF"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        content = []
        documents = self._get_documents(doc_ids)
        
        for i, document in enumerate(documents):
            content.append(Paragraph(document.get('title', 'Untitled'), styles['Heading1']))
            content.append(Spacer(1, 12))
            if document.get('background'):
                content.append(Paragraph("<b>Background:</b>", styles['Heading2']))
                content.append(Paragraph(document['background'][:1000], styles['Normal']))
            if document.get('scope'):
                content.append(Paragraph("<b>Scope:</b>", styles['Heading2']))
                content.append(Paragraph(document['scope'][:1000], styles['Normal']))
            if include_summary:
                content.append(Paragraph("<b>AI Summary:</b>", styles['Heading2']))
                content.append(Paragraph("[Summary placeholder for future AI integration]", styles['Italic']))
            if i < len(documents) - 1:
                content.append(PageBreak())
        
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_docx(self, doc_ids: List[str], include_summary: bool = False) -> bytes:
        """Export documents to DOCX"""
        doc = DocxDocument()
        documents = self._get_documents(doc_ids)
        
        for i, document in enumerate(documents):
            doc.add_heading(document.get('title', 'Untitled'), level=1)
            if document.get('background'):
                doc.add_heading('Background:', level=2)
                doc.add_paragraph(document['background'][:1000])
            if document.get('scope'):
                doc.add_heading('Scope:', level=2)
                doc.add_paragraph(document['scope'][:1000])
            if include_summary:
                doc.add_heading('AI Summary:', level=2)
                p = doc.add_paragraph('[Summary placeholder for future AI integration]')
                p.italic = True
            if i < len(documents) - 1:
                doc.add_page_break()
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def export_to_csv(self, doc_ids: List[str]) -> bytes:
        """Export documents to CSV"""
        buffer = io.StringIO()
        documents = self._get_documents(doc_ids)
        fieldnames = ['Title', 'File Path', 'Background', 'Scope', 'Headings']
        writer = csv.DictWriter(buffer, fieldnames=fieldnames)
        writer.writeheader()
        for document in documents:
            writer.writerow({
                'Title': document.get('title', ''),
                'File Path': document.get('file_path', ''),
                'Background': document.get('background', '')[:500],
                'Scope': document.get('scope', '')[:500],
                'Headings': ' | '.join(document.get('headings', [])[:5])
            })
        buffer.seek(0)
        return buffer.getvalue().encode('utf-8')
